import docx, os

filePath = os.getcwd() + "/" + "demo.docx"
document = docx.Document(filePath)
paras = document.paragraphs
for i in range(len(paras)):
    para = paras[i]
    # para.style = "Title"
    runs = para.runs
    para.add_run("This is a new run")
    for i in range(len(runs)):
        run = runs[i]
        print(run.text)
        # print(run.bold)
        # print(run.italic)
        # print(run.underline)
        run.underline = True

document.add_paragraph("This is new dynamic paragraph")

outputFile = os.getcwd() + "/" + "demo_output.docx"
document.save(outputFile)
